"""
File parsers. Their job: produce a list[RawPage] (page number + text) that
preserves both prose AND tables — critical for a Price Book spec, where the
actual business rules (prices, discount tiers) live inside tables, not
paragraphs. Tables are converted to Markdown and interleaved into the text
stream in document order, so downstream heading-based chunking treats a
table exactly like any other line of content under its enclosing section.

Structural intelligence beyond that (heading detection, hierarchy) lives in
chunker.py, not here — keeps parsing swappable without touching chunking.
"""
from __future__ import annotations

from pathlib import Path

from app.domain.exceptions.errors import DocumentParsingError
from app.infrastructure.parsing.chunker import RawPage
from app.infrastructure.parsing.table_formatter import table_rows_to_markdown


def _docx_heading_prefix(style_name: str) -> str | None:
    """Map Word paragraph styles to markdown heading markers for the chunker."""
    if not style_name:
        return None
    normalized = style_name.strip().lower()
    if normalized.startswith("heading"):
        digits = "".join(ch for ch in style_name if ch.isdigit())
        level = min(int(digits), 6) if digits else 1
        return "#" * level + " "
    if normalized in {"title", "subtitle"}:
        return "# "
    return None


def _docx_bold_heading_prefix(paragraph) -> str | None:
    """
    Many real-world DOCX files (including the challenge brief) use bold body
    text for section titles instead of Word Heading styles.
    """
    runs = [r for r in paragraph.runs if r.text and r.text.strip()]
    if not runs or not all(r.bold for r in runs):
        return None
    text = paragraph.text.strip()
    if len(text) < 3 or len(text) > 60:
        return None
    if text.endswith(":"):
        return None
    if "." in text or "!" in text or "?" in text:
        return None
    if text.startswith("|") or "{" in text:
        return None
    return "# "


def _numbered_table_lines(rows: list[list[str]]) -> list[str] | None:
    """
    Requirement-style tables with numbered first column (1, 2, 3…) in a
    two-column layout become section headings. Multi-column data tables
    (validation rules, specs) stay as markdown tables.
    """
    cleaned_rows = [[str(c).strip() for c in row] for row in rows if any(str(c).strip() for c in row)]
    if not cleaned_rows:
        return None

    max_cols = max(len(row) for row in cleaned_rows)
    if max_cols > 2:
        return None

    header_text = " ".join(cleaned_rows[0]).lower()
    if any(kw in header_text for kw in ("validation", "component", "field", "action", "checked", "required")):
        return None

    numbered_rows = [r for r in cleaned_rows if r and r[0].isdigit()]
    if len(numbered_rows) < 2:
        return None

    lines: list[str] = []
    for row in cleaned_rows:
        if not row or not any(cell for cell in row):
            continue
        if row[0].isdigit():
            body = " ".join(cell for cell in row[1:] if cell).strip()
            if body:
                title_hint = " ".join(body.split()[:8])
                lines.append(f"## {row[0]} {title_hint}")
                lines.append(body)
            continue
        joined = " | ".join(cell for cell in row if cell)
        if joined:
            lines.append(joined)
    return lines or None


def parse_pdf(file_path: str | Path) -> list[RawPage]:
    """
    Uses pdfplumber (not pypdf) specifically because it exposes both
    `.extract_text()` and `.extract_tables()` per page from the same parse,
    which is what lets us preserve tables as structured Markdown instead of
    the flattened, column-scrambled text pypdf-style text extraction alone
    would produce for a pricing table.
    """
    try:
        import pdfplumber
    except ImportError as e:
        raise DocumentParsingError("pdfplumber is required to parse PDF files") from e

    try:
        pages: list[RawPage] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""

                table_blocks = []
                for raw_table in page.extract_tables() or []:
                    md = table_rows_to_markdown(raw_table)
                    if md:
                        table_blocks.append(md)

                combined = text
                if table_blocks:
                    combined = (combined + "\n\n" + "\n\n".join(table_blocks)).strip()

                pages.append(RawPage(page_number=i, text=combined))

        if not pages:
            raise DocumentParsingError(f"No extractable pages in {file_path}")
        # Reject fully empty extracts (image-only / encrypted / corrupt PDFs).
        if all(not (p.text or "").strip() for p in pages):
            raise DocumentParsingError(
                f"PDF has no extractable text (image-only, encrypted, or corrupt): {file_path}"
            )
        return pages
    except DocumentParsingError:
        raise
    except FileNotFoundError as e:
        raise DocumentParsingError(f"PDF file not found: {file_path}") from e
    except Exception as e:
        raise DocumentParsingError(f"Failed to parse PDF {file_path}: {e}") from e


def _iter_docx_block_items(document):
    """
    python-docx doesn't expose a built-in "give me paragraphs and tables in
    document order" API — its .paragraphs and .tables collections are
    separate and lose relative ordering. This walks the underlying XML body
    directly so a table that appears between two paragraphs stays between
    them in the extracted text, which matters for keeping a pricing table
    under the correct heading.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def parse_docx(file_path: str | Path) -> list[RawPage]:
    try:
        import docx
    except ImportError as e:
        raise DocumentParsingError("python-docx is required to parse DOCX files") from e

    try:
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = docx.Document(str(file_path))
        lines: list[str] = []

        for block in _iter_docx_block_items(document):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text:
                    prefix = _docx_heading_prefix(getattr(block.style, "name", "") or "")
                    if not prefix:
                        prefix = _docx_bold_heading_prefix(block)
                    lines.append(f"{prefix}{text}" if prefix else text)
            elif isinstance(block, Table):
                rows = [[cell.text for cell in row.cells] for row in block.rows]
                numbered = _numbered_table_lines(rows)
                if numbered:
                    lines.extend(numbered)
                else:
                    md = table_rows_to_markdown(rows)
                    if md:
                        lines.append(md)

        full_text = "\n".join(lines).strip()
        if not full_text:
            raise DocumentParsingError(
                f"DOCX has no extractable text content: {file_path}"
            )

        # DOCX has no native page-boundary concept in the object model, so
        # we treat the whole document as one logical page stream and rely
        # on heading styles/numbering for hierarchy. This is called out
        # explicitly in README as a documented limitation for DOCX citations
        # (section is authoritative; page is best-effort/approximate).
        return [RawPage(page_number=1, text=full_text)]
    except DocumentParsingError:
        raise
    except FileNotFoundError as e:
        raise DocumentParsingError(f"DOCX file not found: {file_path}") from e
    except Exception as e:
        raise DocumentParsingError(f"Failed to parse DOCX {file_path}: {e}") from e


def parse_document(file_path: str | Path) -> list[RawPage]:
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix in (".docx", ".dotx"):
        return parse_docx(path)
    raise DocumentParsingError(f"Unsupported file type: {suffix}")
