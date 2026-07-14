"""Parser edge-case tests."""
from __future__ import annotations

from pathlib import Path

import pytest

from app.domain.exceptions.errors import DocumentParsingError
from app.infrastructure.parsing.parsers import parse_document, parse_docx, parse_pdf


def test_unsupported_extension_raises(tmp_path: Path):
    bad = tmp_path / "notes.txt"
    bad.write_text("hello", encoding="utf-8")
    with pytest.raises(DocumentParsingError, match="Unsupported"):
        parse_document(bad)


def test_missing_pdf_raises(tmp_path: Path):
    with pytest.raises(DocumentParsingError, match="not found|Failed to parse PDF"):
        parse_pdf(tmp_path / "does-not-exist.pdf")


def test_challenge_style_docx_produces_multiple_sections(tmp_path: Path):
    """Regression: brief-style DOCX with bold titles + numbered 2-col tables."""
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_paragraph("Intro paragraph about the challenge.")
    title = doc.add_paragraph("What to Build")
    for run in title.runs:
        run.bold = True
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "1"
    table.rows[0].cells[1].text = "Document Comparison Engine with citations."
    table.rows[1].cells[0].text = "2"
    table.rows[1].cells[1].text = "RAG Chatbot for single document mode."
    table.rows[2].cells[0].text = "3"
    table.rows[2].cells[1].text = "Cross-document chat with citations."

    path = tmp_path / "challenge_style.docx"
    doc.save(path)

    from app.infrastructure.parsing.parsers import parse_docx
    from app.infrastructure.parsing.chunker import build_parent_child_chunks
    from app.domain.entities.document import ChunkType

    pages = parse_docx(path)
    chunks = build_parent_child_chunks(pages, document_id="x", document_name="brief.docx", version="v1")
    parents = [c for c in chunks if c.chunk_type == ChunkType.PARENT]
    assert len(parents) >= 4
    assert len(chunks) >= len(parents)


def test_empty_docx_raises(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document

    path = tmp_path / "empty.docx"
    Document().save(path)
    with pytest.raises(DocumentParsingError, match="no extractable text"):
        parse_docx(path)
